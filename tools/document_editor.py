"""Édition de documents longs (.docx, ex. romans) avec MODIFICATIONS SUIVIES Word.

Principe (cf. DEV_NOTES) : on ne TOUCHE JAMAIS l'original. On télécharge le .docx depuis
Nextcloud vers un workspace DÉDIÉ, l'agent le révise CHAPITRE PAR CHAPITRE, et chaque
révision est écrite comme des **modifications suivies** (`w:ins`/`w:del`, auteur « Athena »).
À la publication, une COPIE « <nom> — révisé.docx » est déposée sur Nextcloud : tu l'ouvres
dans OnlyOffice, tu vois les ajouts/suppressions et tu les acceptes/refuses une par une.

Dépendance : python-docx. Téléchargement/upload via core.nextcloud (WebDAV).
"""
import os
import re
import json
import difflib
import datetime
import urllib.parse

import requests

from core import nextcloud, projects, user_config
from tools.net_guard import is_blocked_url

_AUTHOR = "Athena"


def _dir() -> str:
    """Dossier de travail DÉDIÉ aux documents en cours d'édition (par utilisateur)."""
    slug = re.sub(r"[^A-Za-z0-9_.-]", "_", user_config.current_user_key()) or "local"
    base = os.environ.get("ACTIVE_WORKSPACE_DIR") or "workspace"
    d = os.path.join(base, "redaction", slug)
    os.makedirs(d, exist_ok=True)
    return d


def _safe_name(name: str) -> str:
    """Nom de fichier local sûr (pas de traversal)."""
    return os.path.basename((name or "").strip()) or "document.docx"


def _local_path(name: str) -> str:
    return os.path.join(_dir(), _safe_name(name))


# --- Nextcloud (binaire) ----------------------------------------------------
def _nc_url(remote_path: str) -> str:
    segs = [s for s in (remote_path or "").strip().lstrip("/").split("/") if s not in ("", ".")]
    if any(s == ".." for s in segs):
        raise ValueError("chemin distant invalide ('..').")
    return nextcloud.files_base() + "/".join(urllib.parse.quote(s) for s in segs)


def _now_iso() -> str:
    return datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


# --- python-docx helpers : modifications suivies -----------------------------
def _docx():
    from docx import Document
    return Document


def _make_text_run(text, deltext=False):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    r = OxmlElement("w:r")
    t = OxmlElement("w:delText" if deltext else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


class _Rev:
    """Compteur d'IDs de révision + fabrique d'éléments w:ins / w:del."""
    def __init__(self):
        self.n = 0
        self.date = _now_iso()

    def _id(self):
        self.n += 1
        return str(self.n)

    def _wrap(self, tag, run):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        el = OxmlElement(tag)
        el.set(qn("w:id"), self._id())
        el.set(qn("w:author"), _AUTHOR)
        el.set(qn("w:date"), self.date)
        el.append(run)
        return el

    def ins(self, text):
        return self._wrap("w:ins", _make_text_run(text, deltext=False))

    def dele(self, text):
        return self._wrap("w:del", _make_text_run(text, deltext=True))


def _tokenize(s):
    """Découpe en jetons mots + espaces (pour un diff fin qui préserve l'espacement)."""
    return re.findall(r"\S+|\s+", s or "")


def _tracked_replace_paragraph(paragraph, new_text, rev: "_Rev"):
    """Révise un paragraphe avec un diff MOT À MOT : seuls les fragments modifiés sont marqués
    (ancien en suppression suivie, nouveau en insertion suivie) ; le reste demeure en texte
    normal. Bien plus lisible dans OnlyOffice qu'un paragraphe entier barré/réinséré."""
    from docx.oxml.ns import qn
    p = paragraph._p
    old = paragraph.text
    for child in list(p):
        if child.tag in (qn("w:r"), qn("w:ins"), qn("w:del")):
            p.remove(child)
    old_tok, new_tok = _tokenize(old), _tokenize(new_text)
    sm = difflib.SequenceMatcher(a=old_tok, b=new_tok, autojunk=False)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        old_seg = "".join(old_tok[i1:i2])
        new_seg = "".join(new_tok[j1:j2])
        if tag == "equal":
            if old_seg:
                p.append(_make_text_run(old_seg))          # texte inchangé (normal)
        elif tag == "delete":
            if old_seg:
                p.append(rev.dele(old_seg))
        elif tag == "insert":
            if new_seg:
                p.append(rev.ins(new_seg))
        else:  # replace
            if old_seg:
                p.append(rev.dele(old_seg))
            if new_seg:
                p.append(rev.ins(new_seg))


def _tracked_delete_paragraph(paragraph, rev: "_Rev"):
    from docx.oxml.ns import qn
    p = paragraph._p
    old = paragraph.text
    for child in list(p):
        if child.tag in (qn("w:r"), qn("w:ins"), qn("w:del")):
            p.remove(child)
    if old:
        p.append(rev.dele(old))


def _insert_tracked_paragraph_after(ref_paragraph, text, rev: "_Rev"):
    """Insère un NOUVEAU paragraphe (texte en insertion suivie) après ref_paragraph."""
    new_p = ref_paragraph.insert_paragraph_before("")  # crée avant…
    # …puis on le déplace juste APRÈS ref (insert_paragraph_before n'a pas d'after natif).
    ref_paragraph._p.addnext(new_p._p)
    new_p._p.append(rev.ins(text))
    return new_p


# --- Détection de chapitres -------------------------------------------------
_CHAP_RX = re.compile(r"^\s*(chapitre|chapter|partie|prologue|épilogue|epilogue)\b", re.IGNORECASE)


def _is_heading(paragraph) -> bool:
    style = (getattr(paragraph.style, "name", "") or "").lower()
    if any(k in style for k in ("heading", "titre", "title")):
        return True
    txt = (paragraph.text or "").strip()
    return bool(txt) and len(txt) < 80 and bool(_CHAP_RX.match(txt))


def _chapters(doc):
    """Renvoie [(titre, i_start, i_end_exclu)] : segments entre titres. Tout avant le 1er
    titre = un segment « (début) »."""
    paras = doc.paragraphs
    heads = [i for i, p in enumerate(paras) if _is_heading(p)]
    chapters = []
    if not heads or heads[0] != 0:
        end = heads[0] if heads else len(paras)
        if end > 0:
            chapters.append(("(début)", 0, end))
    for j, h in enumerate(heads):
        end = heads[j + 1] if j + 1 < len(heads) else len(paras)
        chapters.append((paras[h].text.strip() or f"Chapitre {j+1}", h, end))
    return chapters


def _find_chapter(doc, chapter):
    """Trouve un chapitre par titre (sous-chaîne, insensible casse) ou index 1-based."""
    chaps = _chapters(doc)
    if not chapter:
        return None
    c = str(chapter).strip().lower()
    if c.isdigit():
        idx = int(c) - 1
        return chaps[idx] if 0 <= idx < len(chaps) else None
    for title, a, b in chaps:
        if c in title.lower():
            return (title, a, b)
    return None


# --- OUTILS exposés à l'agent ----------------------------------------------
def document_open(nextcloud_path: str) -> str:
    """
    Ouvre un document .docx depuis Nextcloud pour édition (le télécharge dans un espace de
    travail dédié). L'ORIGINAL sur Nextcloud n'est jamais modifié.

    Args:
        nextcloud_path (str): Chemin du .docx sur Nextcloud (ex: "Romans/MonRoman.docx").

    Returns:
        str: Confirmation + liste des chapitres détectés.
    """
    if not nextcloud.is_configured():
        return "Nextcloud non configuré (Réglages → Agenda → section Nextcloud)."
    if not (nextcloud_path or "").lower().endswith(".docx"):
        return "Seuls les fichiers .docx sont pris en charge pour l'édition suivie."
    try:
        url = _nc_url(nextcloud_path)
    except ValueError as e:
        return f"Erreur : {e}"
    if is_blocked_url(url):
        return "Hôte Nextcloud bloqué (anti-SSRF) — ajoute-le à NET_GUARD_ALLOW_HOSTS."
    try:
        r = requests.get(url, auth=nextcloud.auth(), timeout=30)
        if r.status_code != 200:
            return f"Téléchargement impossible ({r.status_code})."
        name = _safe_name(nextcloud_path)
        local = _local_path(name)
        with open(local, "wb") as f:
            f.write(r.content)
        # mémorise le chemin distant d'origine (pour publier à côté)
        with open(local + ".src", "w", encoding="utf-8") as f:
            f.write(nextcloud_path)
        doc = _docx()(local)
        chaps = _chapters(doc)
        lst = "\n".join(f"  {i+1}. {t}  ({b-a} paragraphe(s))" for i, (t, a, b) in enumerate(chaps))
        return (f"📄 « {name} » ouvert pour édition (copie de travail, original intact).\n"
                f"{len(chaps)} chapitre(s) :\n{lst or '  (aucun chapitre détecté)'}")
    except Exception as e:
        return f"Erreur à l'ouverture : {e}"


def document_read(filename: str, chapter: str = "") -> str:
    """
    Lit le contenu d'un document ouvert (tout, ou un chapitre précis).

    Args:
        filename (str): Nom du document ouvert (ex: "MonRoman.docx").
        chapter (str): Titre (ou numéro) du chapitre à lire. Vide = tout le document.

    Returns:
        str: Le texte demandé.
    """
    local = _local_path(filename)
    if not os.path.exists(local):
        return "Document non ouvert. Utilise d'abord document_open(chemin_nextcloud)."
    try:
        doc = _docx()(local)
        paras = doc.paragraphs
        if chapter:
            ch = _find_chapter(doc, chapter)
            if not ch:
                return f"Chapitre '{chapter}' introuvable. document_read sans chapitre pour la liste."
            title, a, b = ch
            body = "\n".join(p.text for p in paras[a:b])
            return f"# {title}\n{body}"
        full = "\n".join(p.text for p in paras) or "(document vide)"
        # Garde-fou CONTEXTE : ne JAMAIS déverser un document entier (un roman = 100k+ car.)
        # dans le contexte → ça sature le modèle et déclenche des hallucinations. Au-delà d'un
        # seuil, on renvoie la liste des chapitres et on invite à lire chapitre par chapitre.
        _CAP = int(os.getenv("DOCUMENT_READ_CAP", "8000") or 8000)
        if len(full) > _CAP:
            chaps = _chapters(doc)
            lst = "\n".join(f"  {i+1}. {t}" for i, (t, a, b) in enumerate(chaps))
            return (f"⚠️ Document volumineux ({len(full)} caractères) — ne lis PAS tout d'un coup "
                    f"(ça sature le contexte). Lis chapitre par chapitre via document_read(\"{filename}\", "
                    f"chapter=\"…\"), ou utilise document_autorevise pour tout réviser d'un coup.\n"
                    f"{len(chaps)} chapitre(s) :\n{lst}")
        return full
    except Exception as e:
        return f"Erreur de lecture : {e}"


def document_revise(filename: str, chapter: str, new_text: str) -> str:
    """
    Applique une révision à UN chapitre en MODIFICATIONS SUIVIES (l'ancien texte est marqué
    supprimé, le nouveau inséré). À répéter chapitre par chapitre. L'original reste intact ;
    seule la copie de travail accumule les révisions.

    Args:
        filename (str): Document ouvert (ex: "MonRoman.docx").
        chapter (str): Titre ou numéro du chapitre à réviser.
        new_text (str): Nouveau texte COMPLET du chapitre (un paragraphe par ligne).

    Returns:
        str: Résumé des modifications appliquées.
    """
    if not projects.can_write():
        return "Erreur : édition non autorisée (accès en lecture seule)."
    local = _local_path(filename)
    if not os.path.exists(local):
        return "Document non ouvert. Utilise d'abord document_open(chemin_nextcloud)."
    try:
        Document = _docx()
        doc = Document(local)
        ch = _find_chapter(doc, chapter)
        if not ch:
            return f"Chapitre '{chapter}' introuvable."
        title, a, b = ch
        paras = doc.paragraphs
        # On ne révise PAS la ligne de titre si c'en est une (a pointe sur le titre).
        body_start = a + 1 if _is_heading(paras[a]) else a
        old_paras = paras[body_start:b]
        old_texts = [p.text for p in old_paras]
        new_lines = [ln for ln in (new_text or "").replace("\r\n", "\n").split("\n") if ln.strip() != "" or True]
        # On retire les lignes vides en tête/queue pour un diff propre.
        while new_lines and new_lines[0].strip() == "":
            new_lines.pop(0)
        while new_lines and new_lines[-1].strip() == "":
            new_lines.pop()

        rev = _Rev()
        sm = difflib.SequenceMatcher(a=old_texts, b=new_lines, autojunk=False)
        ins_c = del_c = repl_c = 0
        # On parcourt les opcodes ; pour insérer, on s'ancre sur le paragraphe précédent.
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == "equal":
                continue
            if tag == "replace":
                k = 0
                while i1 + k < i2 and j1 + k < j2:
                    _tracked_replace_paragraph(old_paras[i1 + k], new_lines[j1 + k], rev)
                    repl_c += 1
                    k += 1
                # surplus d'anciens → suppression suivie
                for p in old_paras[i1 + k:i2]:
                    _tracked_delete_paragraph(p, rev)
                    del_c += 1
                # surplus de nouveaux → insertion après le dernier paragraphe traité
                anchor = old_paras[min(i2, len(old_paras)) - 1] if old_paras else None
                for line in new_lines[j1 + k:j2]:
                    if anchor is not None:
                        anchor = _insert_tracked_paragraph_after(anchor, line, rev)
                        ins_c += 1
            elif tag == "delete":
                for p in old_paras[i1:i2]:
                    _tracked_delete_paragraph(p, rev)
                    del_c += 1
            elif tag == "insert":
                anchor = old_paras[i1 - 1] if i1 > 0 and old_paras else (old_paras[0] if old_paras else None)
                for line in new_lines[j1:j2]:
                    if anchor is not None:
                        anchor = _insert_tracked_paragraph_after(anchor, line, rev)
                        ins_c += 1
        doc.save(local)
        return (f"✅ Chapitre « {title} » révisé en modifications suivies : "
                f"{repl_c} paragraphe(s) modifié(s), {ins_c} ajouté(s), {del_c} supprimé(s). "
                f"Utilise document_publish('{filename}') quand tu as fini.")
    except Exception as e:
        return f"Erreur lors de la révision : {e}"


def _llm_corrections(model: str, instruction: str, chapter_text: str) -> list:
    """Demande au LLM la LISTE des corrections ponctuelles (et non une réécriture) au format
    JSON [{"old": "<extrait exact>", "new": "<corrigé>"}]. Bornage par construction : on
    n'appliquera QUE ces fragments → impossible de « changer l'histoire ». Renvoie [] si rien."""
    try:
        from core.state import swarm as _sw
        sys_p = (
            "Tu es un correcteur. On te donne un extrait de roman. Tu renvoies la LISTE des "
            "corrections PONCTUELLES à y apporter, au format JSON STRICT : une liste d'objets "
            "{\"old\": \"...\", \"new\": \"...\"}.\n"
            "RÈGLES :\n"
            "- `old` = un fragment COURT (quelques mots) copié EXACTEMENT du texte (à l'identique, "
            "mêmes accents/ponctuation).\n"
            "- Corrige UNIQUEMENT : fautes d'orthographe, de grammaire, de ponctuation, et lourdeurs "
            "ÉVIDENTES (répétitions, adjectifs/adverbes redondants).\n"
            "- NE reformule PAS des phrases entières, NE change PAS l'histoire, les noms, le sens. "
            "NE crée pas d'entrée si rien n'est fautif.\n"
            "- Réponds UNIQUEMENT par le tableau JSON (commence par [ et finis par ]). Si rien à "
            "corriger : [].")
        usr = (f"Consigne : {instruction or 'corrige les fautes et les lourdeurs évidentes, sans réécrire'}\n\n"
               f"--- EXTRAIT ---\n{chapter_text}")
        resp = _sw._complete(model, [{"role": "system", "content": sys_p},
                                     {"role": "user", "content": usr}], tools_schema=None)
        raw = (resp.choices[0].message.content or "").strip()
        # Extrait le tableau JSON même s'il est entouré de texte / d'un bloc ```json.
        m = re.search(r"\[.*\]", raw, re.DOTALL)
        if not m:
            return []
        data = json.loads(m.group(0))
        out = []
        for d in data:
            if isinstance(d, dict) and (d.get("old") or "").strip() and "new" in d:
                if d["old"] != d["new"]:
                    out.append({"old": str(d["old"]), "new": str(d["new"])})
        return out
    except Exception as e:
        print(f"[document_autorevise] corrections LLM échec : {e}")
        return []


def _apply_corrections_to_chapter(doc, chapter_title, corrections, rev) -> int:
    """Applique les corrections {old→new} aux paragraphes du chapitre, en MODIFICATIONS SUIVIES
    mot à mot (seuls les fragments corrigés sont marqués). Renvoie le nb de paragraphes modifiés."""
    ch = _find_chapter(doc, chapter_title)
    if not ch:
        return 0
    title, a, b = ch
    paras = doc.paragraphs
    body_start = a + 1 if _is_heading(paras[a]) else a
    changed = 0
    for p in paras[body_start:b]:
        original = p.text
        if not original.strip():
            continue
        revised = original
        for corr in corrections:
            if corr["old"] in revised:
                revised = revised.replace(corr["old"], corr["new"], 1)
        if revised != original:
            _tracked_replace_paragraph(p, revised, rev)
            changed += 1
    return changed


def document_autorevise(nextcloud_path: str, instruction: str = "", chapter: str = "") -> str:
    """
    Révise un document .docx Nextcloud DE BOUT EN BOUT en un seul appel : télécharge l'original
    (intact), révise chaque chapitre via le LLM en MODIFICATIONS SUIVIES, puis publie la copie
    « <nom> — révisé.docx » sur Nextcloud. Idéal pour réviser un roman entier sans saturer le
    contexte (chaque chapitre est traité isolément).

    Args:
        nextcloud_path (str): Chemin du .docx sur Nextcloud (ex: "roman/MonRoman.docx").
        instruction (str): Consigne de révision (style, ton…). Optionnel.
        chapter (str): Pour ne réviser QU'UN chapitre (titre ou numéro). Vide = tout le document.

    Returns:
        str: Bilan (chapitres révisés) + résultat de la publication.
    """
    if not projects.can_write():
        return "Erreur : édition non autorisée (accès en lecture seule)."
    res = document_open(nextcloud_path)
    if "📄" not in res:
        return res  # erreur d'ouverture (déjà explicite)
    name = _safe_name(nextcloud_path)
    local = _local_path(name)
    try:
        Document = _docx()
        doc = Document(local)
        chaps = _chapters(doc)
        if chapter:
            ch = _find_chapter(doc, chapter)
            if not ch:
                return f"Chapitre '{chapter}' introuvable."
            targets = [ch[0]]
        else:
            targets = [t for (t, a, b) in chaps]
        if not targets:
            return "Aucun chapitre détecté à réviser."

        from core.state import swarm as _sw
        model = getattr(_sw.agents.get(getattr(_sw, "orchestrator_name", "Athena")), "model", None) or "gpt-4o-mini"

        rev = _Rev()
        done, total_corr = [], 0
        for title in targets:
            old = document_read(name, chapter=title)
            old_body = "\n".join(old.split("\n")[1:]) if old.startswith("# ") else old
            if not old_body.strip():
                continue
            # On demande au LLM la LISTE des corrections ponctuelles (pas une réécriture), puis
            # on n'applique QUE ces fragments → révision fidèle et fine (phrase/mot), jamais le
            # chapitre entier barré.
            corrections = _llm_corrections(model, instruction, old_body)
            if not corrections:
                continue
            n = _apply_corrections_to_chapter(doc, title, corrections, rev)
            if n:
                done.append(f"{title} ({n}¶)")
                total_corr += len(corrections)
        if not done:
            return ("Aucune correction proposée par le modèle (texte déjà propre, ou modèle peu "
                    "coopératif). Rien n'a été modifié.")
        doc.save(local)
        pub = document_publish(name)
        return (f"✅ Révision terminée : {total_corr} correction(s) ponctuelle(s) sur "
                f"{len(done)} chapitre(s) [{', '.join(done)}], en modifications suivies.\n{pub}")
    except Exception as e:
        return f"Erreur lors de la révision automatique : {e}"


def _llm_coherence(model: str, canon: str, chapter_title: str, chapter_text: str) -> dict:
    """Analyse un chapitre vs la « bible » accumulée. Renvoie {"incoherences":[...], "canon": "..."}.
    Contexte borné : seulement la bible (compacte) + le chapitre courant."""
    try:
        from core.state import swarm as _sw
        sys_p = (
            "Tu es un éditeur qui vérifie la COHÉRENCE NARRATIVE d'un roman, chapitre par chapitre. "
            "On te donne la BIBLE (faits établis dans les chapitres précédents) et le chapitre courant. "
            "Renvoie un JSON STRICT : {\"incoherences\": [\"...\"], \"canon\": \"...\"}.\n"
            "- `incoherences` : liste des CONTRADICTIONS du chapitre avec la bible — traits physiques "
            "(yeux, cheveux…), noms/orthographe des personnages et lieux, règles de l'univers/magie, "
            "chronologie, faits déjà établis. Sois précis (cite l'élément). [] si aucune.\n"
            "- `canon` : la bible MISE À JOUR et COMPACTE (≤ 1500 caractères) : personnages clés et "
            "leurs attributs, lieux, règles, faits majeurs. Fusionne l'ancienne bible + le nouveau "
            "chapitre, sans tout recopier.\n"
            "Réponds UNIQUEMENT par le JSON.")
        usr = f"BIBLE ACTUELLE :\n{canon or '(vide — premier chapitre)'}\n\n--- CHAPITRE « {chapter_title} » ---\n{chapter_text}"
        resp = _sw._complete(model, [{"role": "system", "content": sys_p},
                                     {"role": "user", "content": usr}], tools_schema=None)
        raw = (resp.choices[0].message.content or "").strip()
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if not m:
            return {"incoherences": [], "canon": canon}
        d = json.loads(m.group(0))
        inc = [str(x) for x in (d.get("incoherences") or []) if str(x).strip()]
        new_canon = str(d.get("canon") or canon)[:2000]
        return {"incoherences": inc, "canon": new_canon}
    except Exception as e:
        print(f"[document_check_coherence] chapitre '{chapter_title}' : {e}")
        return {"incoherences": [], "canon": canon}


def document_check_coherence(nextcloud_path: str, chapter: str = "") -> str:
    """
    Vérifie la COHÉRENCE NARRATIVE d'un .docx (noms, traits physiques, lieux, règles de l'univers,
    chronologie) chapitre par chapitre, et renvoie un RAPPORT des incohérences détectées.
    LECTURE SEULE : ne modifie pas le document.

    Args:
        nextcloud_path (str): Chemin du .docx sur Nextcloud (ex: "roman/MonRoman.docx").
        chapter (str): Pour n'analyser qu'un chapitre (sinon tout le document).

    Returns:
        str: Rapport de cohérence (incohérences par chapitre, ou « aucune détectée »).
    """
    res = document_open(nextcloud_path)
    if "📄" not in res:
        return res
    name = _safe_name(nextcloud_path)
    try:
        doc = _docx()(_local_path(name))
        chaps = _chapters(doc)
        if chapter:
            ch = _find_chapter(doc, chapter)
            if not ch:
                return f"Chapitre '{chapter}' introuvable."
            targets = [ch[0]]
        else:
            targets = [t for (t, a, b) in chaps]
        from core.state import swarm as _sw
        model = getattr(_sw.agents.get(getattr(_sw, "orchestrator_name", "Athena")), "model", None) or "gpt-4o-mini"

        canon = ""
        report = []
        for title in targets:
            txt = document_read(name, chapter=title)
            body = "\n".join(txt.split("\n")[1:]) if txt.startswith("# ") else txt
            if not body.strip():
                continue
            r = _llm_coherence(model, canon, title, body)
            canon = r["canon"]
            if r["incoherences"]:
                report.append(f"\n📍 {title} :\n" + "\n".join(f"   • {i}" for i in r["incoherences"]))

        if not report:
            return f"✅ Aucune incohérence narrative détectée sur {len(targets)} chapitre(s)."
        return ("🔎 RAPPORT DE COHÉRENCE (vérifie/corrige toi-même — rien n'a été modifié) :\n"
                + "".join(report))
    except Exception as e:
        return f"Erreur lors de la vérification de cohérence : {e}"


def _heading_level(paragraph) -> int:
    """Niveau de titre (1 par défaut) d'après le style 'Heading N' / 'Titre N'."""
    m = re.search(r"(\d+)", getattr(paragraph.style, "name", "") or "")
    return int(m.group(1)) if m else 1


def _translate_block(model: str, lang: str, instruction: str, text: str) -> str:
    """Traduit un bloc en `lang` — traduction VIVANTE (pas littérale). Préserve les paragraphes."""
    try:
        from core.state import swarm as _sw
        sys_p = (
            f"Tu es un traducteur littéraire chevronné. Traduis le passage de roman en {lang}.\n"
            "EXIGENCES :\n"
            f"- Traduction VIVANTE et NATURELLE, JAMAIS littérale : le texte doit se lire comme s'il "
            f"avait été ÉCRIT directement en {lang}.\n"
            "- Préserve la VOIX de l'auteur, le ton, le rythme, le registre et les intentions.\n"
            "- Adapte IDIOMATIQUEMENT les expressions, jeux de mots et tournures (équivalent culturel), "
            "sans calquer la structure de la langue source.\n"
            "- GARDE les noms propres (personnages, lieux) inchangés, sauf équivalent consacré.\n"
            "- Conserve le même DÉCOUPAGE en paragraphes (un paragraphe par ligne).\n"
            "- Réponds UNIQUEMENT par la traduction, sans note, sans commentaire, sans guillemets.")
        usr = (f"{('Consigne : ' + instruction) if instruction else ''}\n\n--- PASSAGE À TRADUIRE ---\n{text}").strip()
        resp = _sw._complete(model, [{"role": "system", "content": sys_p},
                                     {"role": "user", "content": usr}], tools_schema=None)
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        print(f"[document_translate] bloc : {e}")
        return ""


def document_translate(nextcloud_path: str, target_language: str, instruction: str = "") -> str:
    """
    Traduit un roman .docx dans une langue cible, de façon VIVANTE (littéraire, naturelle, pas
    mot-à-mot), chapitre par chapitre. Crée un NOUVEAU fichier « <nom> (<langue>).docx » sur
    Nextcloud — l'original reste intact.

    Args:
        nextcloud_path (str): Chemin du .docx source sur Nextcloud (ex: "roman/MonRoman.docx").
        target_language (str): Langue cible (ex: "anglais", "espagnol", "english").
        instruction (str): Consigne de style optionnelle (ton, public…).

    Returns:
        str: Confirmation + nom du fichier traduit publié.
    """
    if not projects.can_write():
        return "Erreur : édition non autorisée (accès en lecture seule)."
    if not (target_language or "").strip():
        return "Précise la langue cible (ex: anglais)."
    res = document_open(nextcloud_path)
    if "📄" not in res:
        return res
    name = _safe_name(nextcloud_path)
    try:
        Document = _docx()
        src = Document(_local_path(name))
        chaps = _chapters(src)
        paras = src.paragraphs
        out = Document()
        _CHUNK = int(os.getenv("DOCUMENT_TRANSLATE_CHUNK", "5000") or 5000)
        n_chunks = 0
        from core.state import swarm as _sw
        model = getattr(_sw.agents.get(getattr(_sw, "orchestrator_name", "Athena")), "model", None) or "gpt-4o-mini"

        for (title, a, b) in chaps:
            seg = paras[a:b]
            idx = 0
            if seg and _is_heading(seg[0]):
                lvl = _heading_level(seg[0])
                tt = _translate_block(model, target_language, "", seg[0].text) or seg[0].text
                out.add_heading(tt.strip().splitlines()[0] if tt.strip() else seg[0].text, level=min(lvl, 9))
                idx = 1
            # Corps : on regroupe les paragraphes en lots bornés (contexte) puis on traduit.
            body = [p.text for p in seg[idx:]]
            batch, blen = [], 0
            def _flush(batch):
                nonlocal n_chunks
                if not batch:
                    return
                src_text = "\n".join(batch)
                tr = _translate_block(model, target_language, instruction, src_text)
                n_chunks += 1
                lines = tr.split("\n") if tr else batch
                for ln in lines:
                    out.add_paragraph(ln)
            for line in body:
                if blen + len(line) > _CHUNK and batch:
                    _flush(batch); batch, blen = [], 0
                batch.append(line); blen += len(line) + 1
            _flush(batch)

        # Publication : nouveau fichier « <nom> (<langue>).docx » à côté de l'original.
        local_out = os.path.join(_dir(), name[:-5] + f" ({target_language}).docx")
        out.save(local_out)
        original_remote = open(_local_path(name) + ".src", encoding="utf-8").read().strip()
        folder = original_remote.rsplit("/", 1)[0] if "/" in original_remote else ""
        remote = (folder + "/" if folder else "") + os.path.basename(local_out)
        url = _nc_url(remote)
        if is_blocked_url(url):
            return "Hôte Nextcloud bloqué (anti-SSRF)."
        with open(local_out, "rb") as f:
            data = f.read()
        r = requests.put(url, auth=nextcloud.auth(), data=data,
                         headers={"Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
                         timeout=30)
        if r.status_code in (200, 201, 204):
            return (f"🌍 Traduction ({target_language}) publiée sur Nextcloud : « {remote} » "
                    f"({len(chaps)} chapitre(s), {n_chunks} bloc(s) traduits). L'original est intact.")
        return f"Traduction faite mais upload échoué ({r.status_code})."
    except Exception as e:
        return f"Erreur lors de la traduction : {e}"


def document_publish(filename: str) -> str:
    """
    Publie la copie révisée sur Nextcloud sous « <nom> — révisé.docx » (à côté de l'original,
    qui reste INTACT). Ouvre ce fichier dans OnlyOffice pour voir/accepter les modifications.

    Args:
        filename (str): Document ouvert (ex: "MonRoman.docx").

    Returns:
        str: Confirmation + nom du fichier publié.
    """
    if not projects.can_write():
        return "Erreur : publication non autorisée (accès en lecture seule)."
    local = _local_path(filename)
    if not os.path.exists(local):
        return "Document non ouvert."
    src_file = local + ".src"
    if not os.path.exists(src_file):
        return "Chemin Nextcloud d'origine inconnu (ré-ouvre via document_open)."
    try:
        original_remote = open(src_file, encoding="utf-8").read().strip()
        folder = original_remote.rsplit("/", 1)[0] if "/" in original_remote else ""
        base = _safe_name(filename)
        revised_name = base[:-5] + " — révisé.docx" if base.lower().endswith(".docx") else base + " — révisé.docx"
        remote = (folder + "/" + revised_name) if folder else revised_name
        url = _nc_url(remote)
        if is_blocked_url(url):
            return "Hôte Nextcloud bloqué (anti-SSRF)."
        with open(local, "rb") as f:
            data = f.read()
        r = requests.put(url, auth=nextcloud.auth(), data=data,
                         headers={"Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
                         timeout=30)
        if r.status_code in (200, 201, 204):
            return (f"📤 Publié sur Nextcloud : « {remote} » (l'original « {original_remote} » est intact). "
                    "Ouvre-le dans OnlyOffice → tu verras les modifications suivies à accepter/refuser.")
        return f"Échec de l'upload ({r.status_code})."
    except Exception as e:
        return f"Erreur de publication : {e}"
