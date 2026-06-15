"""Édition de documents .docx en modifications suivies : révision d'un chapitre, original
préservé, gardes (lecture seule / non ouvert). Réseau Nextcloud non sollicité (fichier local)."""
import os
import sys
import zipfile
from unittest import mock

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import tools.document_editor as de  # noqa: E402

try:
    import docx  # noqa: F401
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False


def _make_doc():
    from docx import Document
    d = de._dir()
    path = os.path.join(d, "RomanTest.docx")
    doc = Document()
    doc.add_heading("Chapitre 1", level=1)
    doc.add_paragraph("Le héros marcha lentement vers la porte.")
    doc.add_paragraph("Il faisait froid.")
    doc.add_heading("Chapitre 2", level=1)
    doc.add_paragraph("Texte intact du chapitre deux.")
    doc.save(path)
    open(path + ".src", "w", encoding="utf-8").write("Romans/RomanTest.docx")
    return path


def test_not_open_is_graceful():
    out = de.document_read("Inexistant.docx")
    assert "non ouvert" in out.lower(), out


def test_revise_creates_tracked_changes_and_keeps_original():
    if not HAS_DOCX:
        print("OK (python-docx absent — test sauté)")
        return
    path = _make_doc()
    out = de.document_revise(
        "RomanTest.docx", "Chapitre 1",
        "Le héros s'élança vers la porte, le cœur battant.\nIl faisait froid.\nUne ombre le suivait.")
    assert "révisé" in out.lower(), out
    xml = zipfile.ZipFile(path).read("word/document.xml").decode("utf-8")
    assert "<w:del " in xml and "<w:ins " in xml, "modifications suivies absentes"
    assert "marcha lentement" in xml, "ancien texte (suppression) absent"
    assert "s'élança" in xml and "Une ombre" in xml, "nouveau texte (insertion) absent"
    # Chapitre 2 NON touché.
    assert "Texte intact du chapitre deux." in xml and "Chapitre 2" in xml


def test_revise_blocked_when_readonly():
    if not HAS_DOCX:
        print("OK (python-docx absent — test sauté)")
        return
    _make_doc()
    with mock.patch.object(de.projects, "can_write", return_value=False):
        out = de.document_revise("RomanTest.docx", "Chapitre 1", "x")
    assert "lecture seule" in out.lower(), out


def test_autorevise_one_call_revises_all_chapters():
    if not HAS_DOCX:
        print("OK (python-docx absent — test sauté)")
        return
    import zipfile
    import core.state as st
    path = _make_doc()

    class _R:
        def __init__(self, c):
            self.choices = [type("C", (), {"message": type("M", (), {"content": c})()})()]

    def fake_complete(model, messages, tools_schema=None, **k):
        old = messages[-1]["content"].split("--- TEXTE À RÉVISER ---")[-1].strip()
        return _R("\n".join(l + " [revu]" for l in old.split("\n") if l.strip()))

    with mock.patch.object(de, "document_open", lambda p: "📄 ouvert"), \
         mock.patch.object(de, "document_publish", lambda f: "📤 publié (mock)"), \
         mock.patch.object(st.swarm, "_complete", fake_complete):
        out = de.document_autorevise("Romans/RomanTest.docx", "améliore le style")
    assert "révisé" in out.lower(), out
    xml = zipfile.ZipFile(path).read("word/document.xml").decode("utf-8")
    assert "<w:ins " in xml and "<w:del " in xml and "[revu]" in xml, "révision auto non appliquée"


def test_read_caps_large_document():
    if not HAS_DOCX:
        print("OK (python-docx absent — test sauté)")
        return
    from docx import Document
    d = de._dir()
    path = os.path.join(d, "Gros.docx")
    doc = Document()
    doc.add_heading("Chapitre 1", level=1)
    for _ in range(400):
        doc.add_paragraph("Phrase de remplissage assez longue pour dépasser le seuil de lecture. " * 3)
    doc.save(path)
    open(path + ".src", "w", encoding="utf-8").write("roman/Gros.docx")
    out = de.document_read("Gros.docx")  # tout le doc → doit être plafonné
    assert "volumineux" in out.lower() and "chapitre par chapitre" in out.lower(), out


if __name__ == "__main__":
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            fn()
            print(f"OK {name}")
    print("\nTous les tests document_editor passent.")
